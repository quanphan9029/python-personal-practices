import random
from docx import Document

document = Document()

def cho_bai2(pheptinh):
    so1 = 1
    so2 = 1
    ketqua = 0

    if pheptinh == "cong":
        so1 = random.randint(1000, 9999)
        so2 = random.randint(1000, 9999)
        document.add_paragraph(f"{so1} + {so2} =")
    elif pheptinh == "tru" or pheptinh == "-":
        so1 = random.randint(1000, 9999)
        so2 = random.randint(1000, 9999)
        while so1 < so2:
            so1 = random.randint(1000, 9999)
            so2 = random.randint(1000, 9999)
        document.add_paragraph(f"{so1} - {so2} = ")
    elif pheptinh == "nhan" or pheptinh == "*":
        so1 = random.randint(100, 9999)
        so2 = random.randint(2, 9)
        document.add_paragraph(f"{so1} * {so2} = ")
    elif pheptinh == "nhan2" or pheptinh == "*2":
        so1 = random.randint(100, 9999)
        so2 = random.randint(2, 9)
        document.add_paragraph(f"{so2} * {so1} = ")
    elif pheptinh == "chia" or pheptinh == "/":
        so1 = random.randint(100, 9999)
        so2 = random.randint(2, 9)        
        while so1 < so2:    
            so1 = random.randint(100, 9999)
            so2 = random.randint(2, 9)
        document.add_paragraph(f"{so1} : {so2} = ")

list_pheptinh = ["cong", "tru", "nhan", "nhan2", "chia"]
count_cong = 0
count_tru = 0
count_nhan = 0
count_nhan2 = 0
count_chia = 0

while list_pheptinh != []:
    pheptinh = random.choice(list_pheptinh)
    if count_cong == 250:
        list_pheptinh.remove("cong")
        print(f"Plus is done. {count_cong} times.")
        count_cong = 0
    elif count_tru == 250:
        list_pheptinh.remove("tru")
        print(f"Minus s done. {count_tru} times.")
        count_tru = 0
    elif count_nhan == 200:
        list_pheptinh.remove("nhan")
        print(f"Time is done. {count_nhan} times.")
        count_nhan = 0
    elif count_nhan2 == 50:
        list_pheptinh.remove("nhan2")
        print(f"Time 2 is done. {count_nhan2} times.")
        count_nhan2 = 0
    elif count_chia == 250:
        list_pheptinh.remove("chia")
        print(f"Division is done. {count_chia} times.")
        count_chia = 0
    elif pheptinh == "cong":
        count_cong+=1
        cho_bai2(pheptinh)
    elif pheptinh == "tru":
        count_tru+=1
        cho_bai2(pheptinh)
    elif pheptinh == "nhan":
        count_nhan+=1
        cho_bai2(pheptinh)
    elif pheptinh == "nhan2":
        count_nhan2+=1
        cho_bai2(pheptinh)
    elif pheptinh == "chia":
        count_chia+=1
        cho_bai2(pheptinh)
    print(list_pheptinh)
print(list_pheptinh)
ten_file = input(f"Enter your file name: ")

document.save(ten_file)