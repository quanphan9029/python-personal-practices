from ntpath import join
from docx import Document
import random

document = Document()

def cho_bai(pheptinh):
    so1 = 1
    so2 = 1
    ketqua = 0
    
    if pheptinh == "cong" or pheptinh == "+":
        so1 = random.randint(1000, 9999)
        so2 = random.randint(1000, 9999)
        document.add_paragraph(f"{so1} + {so2} =")
    elif pheptinh == "tru" or pheptinh == "-":
        so1 = random.randint(1000, 9999)
        so2 = random.randint(1000, 9999)
        while so1 < so2:
            so1 = random.randint(1000, 9999)
            so2 = random.randint(1000, 9999)
        document.add_paragraph(f"{so1} - {so2} = ")
    elif pheptinh == "nhan" or pheptinh == "*":
        so1 = random.randint(100, 9999)
        so2 = random.randint(2, 9)
        document.add_paragraph(f"{so1} * {so2} = ")
    elif pheptinh == "nhan2" or pheptinh == "*2":
        so1 = random.randint(100, 9999)
        so2 = random.randint(2, 9)
        document.add_paragraph(f"{so2} * {so1} = ")
    elif pheptinh == "chia" or pheptinh == "/":
        so1 = random.randint(100, 9999)
        so2 = random.randint(2, 9)        
        while so1 < so2:    
            so1 = random.randint(100, 9999)
            so2 = random.randint(2, 9)
        document.add_paragraph(f"{so1} : {so2} = ")


list_pheptinh = ["cong", "tru", "nhan", "nhan2", "chia"]

for i in range(0, len(list_pheptinh)):
    pheptinh = list_pheptinh[i]
    print(pheptinh)
    if pheptinh == "nhan":
        sobai = 0
        while sobai in range(0, 200):
            cho_bai(pheptinh)
            sobai+=1
    elif pheptinh == "nhan2":
        sobai = 0
        while sobai in range(0, 50):
            cho_bai(pheptinh)
            sobai+=1
    else:
        sobai = 0
        while sobai in range(0, 250):
            cho_bai(pheptinh)
            sobai+=1

ten_bai = input(f"Enter the file name: ")

document.save(ten_bai)